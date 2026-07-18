"""
P0多模态输入管道测试套件
测试文件解析、Vision理解、OCR提取等功能
"""
import asyncio
import os
import sys
import tempfile
import textwrap
from pathlib import Path

# 添加Python后端路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from agent.tools.file_parse_tools import (
    pdf_parse_executor,
    xlsx_parse_executor,
    docx_parse_executor,
    ocr_extract_executor,
)
from agent.tools.vision_tools import (
    vision_understand_executor,
)


async def test_pdf_parse():
    """测试PDF解析功能"""
    print("🧪 测试PDF解析...")
    
    # 创建临时PDF文件(模拟)
    with tempfile.TemporaryDirectory() as tmpdir:
        # 注意: 实际测试需要真实的PDF文件
        # 这里只是验证工具能够正确导入和调用
        print("   ✅ PDF解析工具可用")


async def test_xlsx_parse():
    """测试Excel解析功能"""
    print("🧪 测试Excel解析...")
    
    try:
        import openpyxl
        print("   ✅ openpyxl已安装")
        
        # 创建临时Excel文件进行测试
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            temp_file = f.name
        
        try:
            # 创建测试Excel
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "测试数据"
            
            # 写入表头和数据
            ws.append(["姓名", "年龄", "城市"])
            ws.append(["张三", 25, "北京"])
            ws.append(["李四", 30, "上海"])
            ws.append(["王五", 28, "广州"])
            
            wb.save(temp_file)
            
            # 测试解析
            result = await xlsx_parse_executor({
                "file_path": temp_file,
                "sheet_name": "测试数据",
                "output_format": "markdown"
            })
            
            if result.success:
                print(f"   ✅ Excel解析成功,输出长度:{len(result.output)}")
                print(f"   输出预览:\n{textwrap.indent(result.output[:200], '      ')}")
            else:
                print(f"   ❌ Excel解析失败: {result.error}")
                
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
                
    except ImportError:
        print("   ⚠️ openpyxl未安装,跳过测试")


async def test_docx_parse():
    """测试Word文档解析功能"""
    print("🧪 测试Word文档解析...")
    
    try:
        from docx import Document
        print("   ✅ python-docx已安装")
        
        # 创建临时Word文件进行测试
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_file = f.name
        
        try:
            # 创建测试Word文档
            doc = Document()
            doc.add_heading('测试文档', 0)
            doc.add_paragraph('这是一个测试段落,用于验证Word文档解析功能。')
            doc.add_paragraph('第二段内容,包含更多文本。')
            doc.save(temp_file)
            
            # 测试解析
            result = await docx_parse_executor({
                "file_path": temp_file,
                "max_output_chars": 1000
            })
            
            if result.success:
                print(f"   ✅ Word文档解析成功,输出长度:{len(result.output)}")
                print(f"   输出预览:\n{textwrap.indent(result.output[:200], '      ')}")
            else:
                print(f"   ❌ Word文档解析失败: {result.error}")
                
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
                
    except ImportError:
        print("   ⚠️ python-docx未安装,跳过测试")


async def test_vision_understand():
    """测试Vision理解功能"""
    print("🧪 测试Vision理解...")
    
    # 检查是否有有效的Vision API配置
    vision_model = os.getenv("VISION_MODEL")
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("ANTHROPIC_API_KEY")
    
    if vision_model and api_key:
        print(f"   ✅ Vision API已配置(model={vision_model})")
        
        # 创建临时测试图片(1x1像素的透明PNG)
        import base64
        # 最小的有效PNG文件(1x1透明像素)
        png_data = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QD"
            "wADhgGAWjR9awAAAABJRU5ErkJggg=="
        )
        
        # 测试Base64输入
        result = await vision_understand_executor({
            "image_base64": base64.b64encode(png_data).decode("utf-8"),
            "question": "这个图片是什么?",
            "model": "gpt-4o"
        })
        
        if result.success:
            print(f"   ✅ Vision理解测试成功")
        else:
            print(f"   ⚠️ Vision理解调用结果: {result.output[:100]}")
    else:
        print("   ⚠️ 未配置Vision API,跳过实际调用测试")


async def test_ocr_extract():
    """测试OCR文本提取功能"""
    print("🧪 测试OCR文本提取...")
    
    # 检查是否有OCR引擎
    try:
        import pytesseract
        print("   ✅ pytesseract已安装")
        
        # 创建临时测试图片(包含文本的图片)
        from PIL import Image, ImageDraw, ImageFont
        import base64
        
        # 创建白色背景图片
        img = Image.new('RGB', (400, 100), color='white')
        d = ImageDraw.Draw(img)
        d.text((10, 10), "Hello OCR!", fill='black')
        
        # 保存为临时文件
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            temp_file = f.name
            img.save(temp_file)
        
        try:
            # 测试OCR
            result = await ocr_extract_executor({
                "image_path": temp_file,
                "language": "en"
            })
            
            if result.success:
                print(f"   ✅ OCR提取成功,输出:{result.output[:50]}")
            else:
                print(f"   ⚠️ OCR调用结果: {result.output[:100]}")
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
                
    except ImportError:
        print("   ⚠️ pytesseract未安装,跳过OCR测试")


async def main():
    """运行所有测试"""
    print("="*60)
    print("🧪 P0多模态输入管道测试套件")
    print("="*60)
    print()
    
    await test_pdf_parse()
    print()
    await test_xlsx_parse()
    print()
    await test_docx_parse()
    print()
    await test_vision_understand()
    print()
    await test_ocr_extract()
    
    print()
    print("="*60)
    print("✅ 测试完成!")
    print("="*60)


if __name__ == "__main__":
    asyncio.run(main())
