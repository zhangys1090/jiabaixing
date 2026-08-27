from __future__ import annotations

import os
import tempfile
from pathlib import Path
from typing import Any
import logging
logger = logging.getLogger(__name__)

try:
    from pdfminer.high_level import extract_text, extract_pages
    from pdfminer.layout import LAParams, LTTextBox
    HAS_PDFMINER = True
except ImportError:
    HAS_PDFMINER = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from agent.tools.registry import (
    ToolCategory,
    ToolDefinition,
    ToolParameterDef,
    ToolResult,
)

# ========== PDF解析工具 ==========

PDF_PARSE_DEF = ToolDefinition(
    name="pdf_parse",
    description="解析PDF文件并提取文本内容。适用场景：阅读PDF文档、提取PDF中的数据。不支持加密PDF。",
    short_desc="解析PDF文件提取文本",
    category=ToolCategory.FILE,
    tags=["file", "pdf", "parse", "extract", "document"],
    scenes=["research", "daily", "development"],
    capability_level=2,
    parameters=[
        ToolParameterDef(name="file_path", type="string", description="PDF文件路径"),
        ToolParameterDef(name="page_range", type="string", required=False, description="要解析的页码范围，如'1-5'或'all'，默认'all'"),
        ToolParameterDef(name="max_output_chars", type="number", required=False, description="最大输出字符数，默认50000"),
    ],
    risk_level="low",
)


# ========== Excel解析工具 ==========

XLSX_PARSE_DEF = ToolDefinition(
    name="xlsx_parse",
    description="解析Excel文件(.xlsx/.xls)并提取数据。适用场景：读取电子表格数据、分析Excel图表。支持多sheet选择。",
    short_desc="解析Excel文件提取数据",
    category=ToolCategory.FILE,
    tags=["file", "excel", "xlsx", "parse", "spreadsheet", "data"],
    scenes=["research", "daily", "finance"],
    capability_level=2,
    parameters=[
        ToolParameterDef(name="file_path", type="string", description="Excel文件路径"),
        ToolParameterDef(name="sheet_name", type="string", required=False, description="要解析的sheet名称，默认第一个sheet"),
        ToolParameterDef(name="header_row", type="number", required=False, description="表头所在行号，默认1"),
        ToolParameterDef(name="max_rows", type="number", required=False, description="最大读取行数，默认1000"),
        ToolParameterDef(name="output_format", type="string", required=False, description="输出格式:csv/json/markdown，默认markdown", enum=["csv", "json", "markdown"]),
    ],
    risk_level="low",
)


# ========== Word文档解析工具 ==========

DOCX_PARSE_DEF = ToolDefinition(
    name="docx_parse",
    description="解析Word文档(.docx)并提取文本内容。适用场景：阅读Word文档、提取文档中的标题和段落。",
    short_desc="解析Word文档提取文本",
    category=ToolCategory.FILE,
    tags=["file", "word", "docx", "parse", "extract", "document"],
    scenes=["research", "daily", "development"],
    capability_level=2,
    parameters=[
        ToolParameterDef(name="file_path", type="string", description="Word文档路径"),
        ToolParameterDef(name="max_output_chars", type="number", required=False, description="最大输出字符数，默认50000"),
    ],
    risk_level="low",
)


# image_understand 已移除 — 与 vision_understand (vision_tools) 重复，统一使用 vision_understand


# ========== OCR文本提取工具 ==========

OCR_EXTRACT_DEF = ToolDefinition(
    name="ocr_extract",
    description="从图片中提取文本(Optical Character Recognition)。适用场景：扫描件文字提取、图片中的文字识别、验证码识别。",
    short_desc="从图片中提取文字",
    category=ToolCategory.FILE,
    tags=["ocr", "image", "text-extract", "scan", "photo"],
    scenes=["research", "daily"],
    capability_level=2,
    parameters=[
        ToolParameterDef(name="image_url", type="string", required=False, description="图片URL或本地路径"),
        ToolParameterDef(name="image_base64", type="string", required=False, description="图片Base64编码数据"),
        ToolParameterDef(name="language", type="string", required=False, description="文字语言，默认中英文混合", enum=["zh", "en", "zh-en"]),
    ],
    risk_level="low",
)

# 全局常量
_MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
_MAX_OUTPUT_CHARS = 50000


# ========== PDF解析执行器 ==========

async def pdf_parse_executor(params: dict[str, Any]) -> ToolResult:
    if not HAS_PDFMINER:
        return ToolResult(success=False, error="pdfminer.skel未安装，请运行: pip install pdfminer.six")

    file_path = str(params.get("file_path", ""))
    page_range = str(params.get("page_range", "all"))
    max_chars = int(params.get("max_output_chars", _MAX_OUTPUT_CHARS))

    if not file_path or not os.path.exists(file_path):
        return ToolResult(success=False, error=f"文件不存在: {file_path}")

    try:
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        if file_size > _MAX_FILE_SIZE:
            return ToolResult(success=False, error=f"文件过大({file_size // 1024 // 1024}MB)")

        # 解析PDF
        laparams = LAParams()

        if page_range.lower() == "all":
            text = extract_text(file_path, laparams=laparams)
        else:
            # 解析页码范围
            start_page, end_page = _parse_page_range(page_range)
            # 获取总页数
            with open(file_path, 'rb') as f:
                from pdfminer.pdfparser import PDFParser
                from pdfminer.pdfdocument import PDFDocument
                from pdfminer.pdfpage import PDFPage
                parser = PDFParser(f)
                document = PDFDocument(parser)
                total_pages = min(len(list(PDFPage.create_pages(document))), 9999)

            # 提取指定页面
            all_pages_text = extract_text(file_path, laparams=laparams)
            pages = all_pages_text.split('\f')  # PDFminer使用换页符分隔页面
            if start_page <= len(pages) and end_page <= len(pages):
                selected_pages = pages[start_page-1:end_page]
                text = '\f'.join(selected_pages)
            else:
                return ToolResult(success=False, error=f"页码超出范围(总共{len(pages)}页)")

        # 截断输出
        if len(text) > max_chars:
            text = text[:max_chars] + "\n...(内容过长已截断)"

        return ToolResult(
            success=True,
            output=text,
            metadata={
                "file_path": file_path,
                "pages_parsed": page_range,
                "total_chars": len(text),
            }
        )
    except Exception as e:
        logger.warning("file_parse_tools 异常处理", error=str(e))
        return ToolResult(success=False, error=f"PDF解析失败: {e}")


# ========== Excel解析执行器 ==========

async def xlsx_parse_executor(params: dict[str, Any]) -> ToolResult:
    if not HAS_OPENPYXL:
        return ToolResult(success=False, error="openpyxl未安装，请运行: pip install openpyxl")

    file_path = str(params.get("file_path", ""))
    sheet_name = str(params.get("sheet_name", ""))  # 空字符串表示第一个sheet
    header_row = int(params.get("header_row", 1))
    max_rows = int(params.get("max_rows", 1000))
    output_format = str(params.get("output_format", "markdown"))

    if not file_path or not os.path.exists(file_path):
        return ToolResult(success=False, error=f"文件不存在: {file_path}")

    try:
        file_size = os.path.getsize(file_path)
        if file_size > _MAX_FILE_SIZE:
            return ToolResult(success=False, error=f"文件过大({file_size // 1024 // 1024}MB)")

        wb = openpyxl.load_workbook(file_path, data_only=True)

        # 选择sheet
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                return ToolResult(success=False, error=f"Sheet '{sheet_name}'不存在，可用sheets: {wb.sheetnames}")
            ws = wb[sheet_name]
        else:
            ws = wb[wb.sheetnames[0]]

        # 读取数据
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return ToolResult(success=True, output="[空表格]", metadata={"file_path": file_path})

        header = rows[header_row-1] if len(rows) >= header_row else None
        data_rows = rows[header_row:max_rows+header_row]

        # 格式化输出
        if output_format == "json":
            import json
            result = [{"col" + str(i+1): cell for i, cell in enumerate(row) if header}
                     for row in data_rows]
            output = json.dumps(result, ensure_ascii=False, indent=2)
        elif output_format == "csv":
            import io
            import csv
            buf = io.StringIO()
            writer = csv.writer(buf)
            if header:
                writer.writerow(header)
            for row in data_rows:
                writer.writerow(list(row))
            output = buf.getvalue()
            buf.close()
        else:  # markdown
            output = _format_as_markdown(header, data_rows)

        if len(output) > _MAX_OUTPUT_CHARS:
            output = output[:_MAX_OUTPUT_CHARS] + "\n...(内容过长已截断)"

        return ToolResult(
            success=True,
            output=output,
            metadata={
                "file_path": file_path,
                "sheet": sheet_name or wb.sheetnames[0],
                "rows_parsed": len(data_rows),
                "columns": len(header) if header else 0,
            }
        )
    except Exception as e:
        logger.warning("file_parse_tools 异常处理", error=str(e))
        return ToolResult(success=False, error=f"Excel解析失败: {e}")


def _format_as_markdown(header, rows):
    """将数据格式化为Markdown表格"""
    if not header:
        return "[无表头]"

    # 转换值为字符串
    header_str = [str(h) if h is not None else "" for h in header]
    separator = "| " + " | ".join(["---"] * len(header_str)) + " |"
    rows_str = []
    for row in rows:
        row_str = " | ".join(str(c) if c is not None else "" for c in row)
        rows_str.append("| " + row_str + " |")

    return "| " + " | ".join(header_str) + " |\n" + separator + "\n" + "\n".join(rows_str[:100])  # 最多100行


# ========== Word文档解析执行器 ==========

async def docx_parse_executor(params: dict[str, Any]) -> ToolResult:
    if not HAS_DOCX:
        return ToolResult(success=False, error="python-docx未安装，请运行: pip install python-docx")

    file_path = str(params.get("file_path", ""))
    max_chars = int(params.get("max_output_chars", _MAX_OUTPUT_CHARS))

    if not file_path or not os.path.exists(file_path):
        return ToolResult(success=False, error=f"文件不存在: {file_path}")

    try:
        file_size = os.path.getsize(file_path)
        if file_size > _MAX_FILE_SIZE:
            return ToolResult(success=False, error=f"文件过大({file_size // 1024 // 1024}MB)")

        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        text = "\n\n".join(paragraphs)

        if not text.strip():
            return ToolResult(success=True, output="[空文档]", metadata={"file_path": file_path})

        # 尝试提取表格
        tables_text = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(" | ".join(cells))
            tables_text.append("\n".join(rows))

        if tables_text:
            text += "\n\n---\n\n## 表格\n\n" + "\n\n---\n\n".join(tables_text)

        if len(text) > max_chars:
            text = text[:max_chars] + "\n...(内容过长已截断)"

        return ToolResult(
            success=True,
            output=text,
            metadata={
                "file_path": file_path,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "total_chars": len(text),
            }
        )
    except Exception as e:
        logger.warning("file_parse_tools 异常处理", error=str(e))
        return ToolResult(success=False, error=f"Word文档解析失败: {e}")


# image_understand_executor 已移除 — 与 vision_understand (vision_tools) 重复


# ========== OCR文本提取执行器 ==========

async def ocr_extract_executor(params: dict[str, Any]) -> ToolResult:
    # 这里可以根据实际部署的OCR引擎实现
    # 例如tesseract、Azure Cognitive Services、百度OCR等
    image_url = params.get("image_url", "")
    image_base64 = params.get("image_base64", "")
    language = params.get("language", "zh-en")

    if not image_url and not image_base64:
        return ToolResult(success=False, error="需要提供图片URL或Base64数据")

    # TODO: 集成实际的OCR API
    # 示例: 使用本地tesseract
    # import pytesseract
    # from PIL import Image
    # if image_base64:
    #     from io import BytesIO
    #     img = Image.open(BytesIO(base64.b64decode(image_base64)))
    #     text = pytesseract.image_to_string(img, lang=language)
    # elif image_url:
    #     # 下载图片后处理
    #     ...

    return ToolResult(
        success=True,
        output="[OCR功能待集成OCR引擎]",
        metadata={
            "has_image_url": bool(image_url),
            "has_base64": bool(image_base64),
            "language": language,
        }
    )


# ========== 辅助函数 ==========

def _parse_page_range(page_range_str: str) -> tuple[int, int]:
    """解析页码范围字符串，如'1-5'或'3'"""
    page_range_str = page_range_str.strip()

    if "-" in page_range_str:
        parts = page_range_str.split("-")
        start = int(parts[0])
        end = int(parts[1]) if len(parts) > 1 else start
        return start, end
    else:
        page_num = int(page_range_str)
        return page_num, page_num
